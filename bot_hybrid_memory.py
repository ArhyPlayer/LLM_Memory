"""
Telegram-бот с гибридной памятью (короткая + долгая)
Объединяет возможности bot_short_memory.py и bot_long_memory.py

Короткая память: последние 10 сообщений диалога (оперативка)
Долгая память: документы в ChromaDB с векторным поиском (RAG)
"""

import os
import json
import logging
import tempfile
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any
from pathlib import Path
from collections import deque

from pydantic import BaseModel
import openai
from aiogram import Bot, Dispatcher, F
from aiogram.filters import Command
from aiogram.types import (
    Message, Document, BotCommand, ReplyKeyboardMarkup, KeyboardButton,
    InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery,
)
from openai import AsyncOpenAI
from dotenv import load_dotenv

import chromadb
from chromadb.config import Settings

# Библиотеки для работы с документами
import PyPDF2
from docx import Document as DocxDocument

# База данных тезисов (SQLite)
import database


# Загрузка переменных окружения
load_dotenv()

# Загрузка настроек
BOT_TOKEN = os.getenv("BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "").strip()
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5-mini")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-3-small")
OPENAI_MAX_COMPLETION_TOKENS = int(os.getenv("OPENAI_MAX_COMPLETION_TOKENS", os.getenv("OPENAI_MAX_TOKENS", "2000")))
OPENAI_TEMPERATURE = float(os.getenv("OPENAI_TEMPERATURE", "0.7"))
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()

# Настройка логирования
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Проверка обязательных параметров
if not BOT_TOKEN:
    raise ValueError("BOT_TOKEN не найден в переменных окружения!")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY не найден в переменных окружения!")

logger.info("🤖 Бот с гибридной памятью инициализирован")
logger.info(f"  📊 Модель: {OPENAI_MODEL}")
logger.info(f"  🧮 Эмбеддинги: {EMBEDDING_MODEL}")
logger.info(f"  📝 Max completion tokens: {OPENAI_MAX_COMPLETION_TOKENS}")
logger.info(f"  🌐 Base URL: {OPENAI_BASE_URL if OPENAI_BASE_URL else 'default'}")

# Инициализация бота и диспетчера
bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

# Инициализация OpenAI клиента
if OPENAI_BASE_URL:
    openai_client = AsyncOpenAI(
        api_key=OPENAI_API_KEY,
        base_url=OPENAI_BASE_URL
    )
else:
    openai_client = AsyncOpenAI(api_key=OPENAI_API_KEY)

# ============================================
# КОРОТКАЯ ПАМЯТЬ (История диалогов)
# ============================================

HISTORY_SIZE = 10  # Размер короткой памяти
user_histories: Dict[int, deque] = {}

def get_user_history(user_id: int) -> deque:
    """Получить историю диалога пользователя"""
    if user_id not in user_histories:
        user_histories[user_id] = deque(maxlen=HISTORY_SIZE)
    return user_histories[user_id]


# ============================================
# РОЛИ ПОВЕДЕНИЯ (модель поведения)
# ============================================

DEFAULT_ROLE = "standard"
ROLES: Dict[str, tuple] = {
    "standard": (
        "Стандартный помощник",
        "Ты — универсальный помощник. Отвечай кратко и по делу, дружелюбно и понятно. "
        "Опирайся на контекст диалога и документы пользователя."
    ),
    "scientific": (
        "Научный",
        "Ты работаешь в научном стиле: точные формулировки, уместные термины, опора на факты и логику. "
        "При необходимости указывай допущения и ограничения. Стиль — академический, без лишней разговорности."
    ),
    "technical": (
        "Технический",
        "Ты — технический ассистент: чёткие пошаговые инструкции, код и команды при необходимости, "
        "минимум воды. Форматируй списки и блоки кода явно. Стиль — документация и гайды."
    ),
    "creative": (
        "Креативный",
        "Ты отвечаешь развёрнуто и образно, предлагаешь идеи и варианты. "
        "Можешь предлагать альтернативы и развивать мысль. Стиль — живой, без сухого перечисления."
    ),
    "concise": (
        "Лаконичный",
        "Ты отвечаешь максимально кратко: только суть, без вступлений и повторов. "
        "Короткие фразы, тезисы, буллеты. Без лишних слов."
    ),
}
user_roles: Dict[int, str] = {}


def get_user_role(user_id: int) -> str:
    """Вернуть ключ роли пользователя (по умолчанию standard)."""
    return user_roles.get(user_id, DEFAULT_ROLE)


def get_role_prompt(role_key: str) -> str:
    """Текст для системного промпта по выбранной роли."""
    name, instruction = ROLES.get(role_key, ROLES[DEFAULT_ROLE])
    return f"\n\nРОЛЬ: {name}.\n{instruction}"


# ============================================
# ДОЛГАЯ ПАМЯТЬ (ChromaDB)
# ============================================

MEMORY_PATH = "./memory"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

Path(MEMORY_PATH).mkdir(exist_ok=True)

chroma_client = chromadb.PersistentClient(
    path=MEMORY_PATH,
    settings=Settings(anonymized_telemetry=False)
)

try:
    collection = chroma_client.get_or_create_collection(
        name="documents",
        metadata={"description": "User documents with embeddings"}
    )
    logger.info(f"📚 ChromaDB готова. Документов: {collection.count()}")
except Exception as e:
    logger.error(f"Ошибка инициализации ChromaDB: {e}")
    raise


# ============================================
# ФУНКЦИИ РАБОТЫ С ДОКУМЕНТАМИ (ДОЛГАЯ ПАМЯТЬ)
# ============================================

def load_document(file_path: str, file_extension: str) -> str:
    """Загрузка и извлечение текста из документа"""
    try:
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            logger.info(f"TXT загружен, длина: {len(text)} символов")
            return text
        
        elif file_extension == '.pdf':
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            logger.info(f"PDF загружен ({len(pdf_reader.pages)} страниц), длина: {len(text)} символов")
            return text
        
        elif file_extension == '.docx':
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            logger.info(f"DOCX загружен ({len(doc.paragraphs)} параграфов), длина: {len(text)} символов")
            return text
        
        else:
            raise ValueError(f"Неподдерживаемый формат: {file_extension}")
    
    except Exception as e:
        logger.error(f"Ошибка загрузки документа: {e}")
        raise


def split_text_into_chunks(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Разбиение текста на чанки с перекрытием"""
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    
    logger.info(f"Текст разбит на {len(chunks)} чанков")
    return chunks


async def embed_chunks(user_id: int, document_name: str, chunks: List[str]) -> int:
    """Создание эмбеддингов и сохранение в ChromaDB"""
    try:
        logger.info(f"Создание эмбеддингов для {len(chunks)} чанков...")
        
        embeddings_response = await openai_client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=chunks
        )
        
        embeddings = [item.embedding for item in embeddings_response.data]
        ids = [f"user_{user_id}_doc_{document_name}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "user_id": str(user_id),
                "document_name": document_name,
                "chunk_index": i,
                "chunk_text": chunks[i][:100]
            }
            for i in range(len(chunks))
        ]
        
        collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas
        )
        
        logger.info(f"✅ Сохранено {len(chunks)} чанков в ChromaDB")
        export_long_memory_to_json()
        return len(chunks)
    
    except Exception as e:
        logger.error(f"Ошибка создания эмбеддингов: {e}")
        raise


async def retrieve_context(user_id: int, query: str, top_k: int = 3) -> List[str]:
    """Поиск релевантных фрагментов в векторной базе"""
    try:
        logger.info(f"Поиск контекста для запроса: '{query[:50]}...'")
        
        query_embedding_response = await openai_client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=[query]
        )
        query_embedding = query_embedding_response.data[0].embedding
        
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            where={"user_id": str(user_id)}
        )
        
        if results and results['documents'] and len(results['documents']) > 0:
            documents = results['documents'][0]
            logger.info(f"Найдено {len(documents)} релевантных фрагментов")
            return documents
        else:
            logger.warning("Релевантные фрагменты не найдены")
            return []
    
    except Exception as e:
        logger.error(f"Ошибка поиска контекста: {e}")
        return []


def get_user_documents_count(user_id: int) -> int:
    """Получить количество чанков пользователя"""
    try:
        results = collection.get(where={"user_id": str(user_id)})
        return len(results['ids']) if results and results['ids'] else 0
    except Exception as e:
        logger.error(f"Ошибка подсчёта документов: {e}")
        return 0


def delete_user_documents(user_id: int) -> int:
    """Удалить все документы пользователя"""
    try:
        results = collection.get(where={"user_id": str(user_id)})
        if results and results['ids']:
            collection.delete(ids=results['ids'])
            logger.info(f"Удалено {len(results['ids'])} чанков для user_id={user_id}")
            count = len(results['ids'])
            export_long_memory_to_json()
            return count
        return 0
    except Exception as e:
        logger.error(f"Ошибка удаления документов: {e}")
        return 0


def export_long_memory_to_json() -> None:
    """
    Сохраняет в memory/memory.json только запросы пользователей —
    текст сообщений, которые они вводят в строке ввода в Telegram.
    """
    try:
        users: Dict[str, Dict[str, Any]] = {}
        for user_id in database.get_all_user_ids_with_requests():
            requests = database.get_all_user_requests(user_id)
            users[str(user_id)] = {"requests": requests}

        data = {
            "updated": datetime.now(timezone.utc).isoformat(),
            "users": users
        }

        memory_json_path = Path(MEMORY_PATH) / "memory.json"
        memory_json_path.parent.mkdir(parents=True, exist_ok=True)
        with open(memory_json_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info("Запросы пользователей сохранены в %s", memory_json_path)
    except Exception as e:
        logger.error("Ошибка экспорта в memory.json: %s", e)


# ============================================
# СТРУКТУРИРОВАННЫЙ ОТВЕТ (тезисы + сообщение)
# ============================================

class DialogueResponse(BaseModel):
    """Структура ответа: тезисы текущего диалога (вопрос + ответ) и сообщение для чата."""
    theses: list[str]
    message: str


# ============================================
# ГИБРИДНАЯ ГЕНЕРАЦИЯ ОТВЕТОВ
# ============================================

async def get_hybrid_response(user_id: int, user_message: str) -> str:
    """
    Генерация ответа с использованием гибридной памяти:
    - Короткая память (история диалога)
    - Долгая память (документы из ChromaDB)
    """
    try:
        # Получаем короткую память (историю диалога)
        history = get_user_history(user_id)
        
        # Проверяем наличие документов в долгой памяти
        has_documents = get_user_documents_count(user_id) > 0
        
        # Формируем базовый список сообщений
        messages = []
        
        # Системный промпт зависит от наличия документов
        structured_instruction = (
            "\n\nВАЖНО: Ты обязан ответить вызовом функции DialogueResponse с двумя полями:\n"
            "1. theses — список кратких тезисов (1–10 пунктов), суммирующих текущий обмен: вопрос пользователя и твой ответ.\n"
            "2. message — текст твоего ответа пользователю (то, что будет показано в чате).\n"
            "Сначала сформулируй ответ в message, затем выдели тезисы диалога в theses."
        )
        # Тезисы из БД (история прошлых диалогов) — подставляем в системный промпт
        db_theses_block = database.get_theses_for_prompt(user_id)
        role_block = get_role_prompt(get_user_role(user_id))

        if has_documents:
            # Ищем релевантный контекст в документах
            context_chunks = await retrieve_context(user_id, user_message, top_k=3)
            
            if context_chunks:
                # Есть релевантные документы - используем их
                context = "\n\n---\n\n".join(context_chunks)
                system_prompt = (
                    "Ты — AI-ассистент с доступом к документам пользователя и истории диалога.\n"
                    + role_block
                    + "\n\nПравила:\n"
                    "1. Используй информацию из документов, когда это релевантно\n"
                    "2. Используй историю диалога для контекста\n"
                    "3. Если информации нет в документах - отвечай на основе диалога\n"
                    "4. Будь полезным, точным и естественным\n\n"
                    f"ДОКУМЕНТЫ ПОЛЬЗОВАТЕЛЯ:\n{context}"
                    + db_theses_block
                    + structured_instruction
                )
            else:
                # Документы есть, но не релевантны к вопросу
                system_prompt = (
                    "Ты — полезный AI-ассистент. У пользователя есть загруженные документы, но они не релевантны к текущему вопросу. Отвечай на основе истории диалога.\n"
                    + role_block
                    + db_theses_block
                    + structured_instruction
                )
        else:
            # Нет документов - обычный диалог
            system_prompt = (
                "Ты — полезный AI-ассистент. Отвечай на основе истории диалога.\n"
                + role_block
                + db_theses_block
                + structured_instruction
            )
        
        messages.append({"role": "system", "content": system_prompt})
        
        # Добавляем историю диалога (короткую память)
        messages.extend(list(history))
        
        # Добавляем текущее сообщение
        messages.append({"role": "user", "content": user_message})
        
        # Параметры запроса для структурированного ответа (parse API)
        api_params = {
            "model": OPENAI_MODEL,
            "messages": messages,
            "tools": [openai.pydantic_function_tool(DialogueResponse)],
            "tool_choice": "required",
            "max_completion_tokens": OPENAI_MAX_COMPLETION_TOKENS
        }
        
        # gpt-5-mini не поддерживает temperature
        if "gpt-5" not in OPENAI_MODEL.lower():
            api_params["temperature"] = OPENAI_TEMPERATURE
        
        logger.info(f"Отправка запроса к API для user_id={user_id}")
        logger.debug(f"Количество сообщений: {len(messages)}, Документы: {has_documents}")
        
        # Запрос к API (structured output: тезисы + сообщение)
        response = await openai_client.beta.chat.completions.parse(**api_params)
        
        tool_calls = response.choices[0].message.tool_calls
        if not tool_calls:
            logger.warning("Нет tool_calls в ответе, используем content как fallback")
            ai_message = response.choices[0].message.content or "Не удалось сформировать ответ."
            history.append({"role": "user", "content": user_message})
            history.append({"role": "assistant", "content": ai_message})
            return ai_message
        
        parsed: DialogueResponse = tool_calls[0].function.parsed_arguments
        ai_message = parsed.message
        theses = parsed.theses
        
        # Тезисы в консоль и в БД (таблица user_<user_id>)
        logger.info("Тезисы диалога: %s", theses)
        database.add_theses(user_id, theses)
        logger.info(f"Ответ получен, длина: {len(ai_message)} символов")
        
        # Сохраняем в короткую память
        history.append({"role": "user", "content": user_message})
        history.append({"role": "assistant", "content": ai_message})
        
        return ai_message
        
    except Exception as e:
        logger.error(f"Ошибка генерации ответа: {type(e).__name__}: {str(e)}")
        
        error_message = "❌ Произошла ошибка при обработке запроса.\n\n"
        
        if "500" in str(e) or "Internal Server Error" in str(e):
            error_message += "🔴 Ошибка сервера ProxyAPI (500).\n"
            error_message += f"Проверьте модель: {OPENAI_MODEL}"
        elif "401" in str(e):
            error_message += "🔑 Ошибка авторизации.\nПроверьте OPENAI_API_KEY"
        elif "429" in str(e):
            error_message += "⏱️ Превышен лимит запросов."
        else:
            error_message += f"Детали: {str(e)}"
        
        return error_message


# ============================================
# ОБРАБОТЧИКИ КОМАНД
# ============================================

# Список команд для меню бота (кнопка рядом со скрепкой)
BOT_COMMANDS = [
    BotCommand(command="start", description="Начать работу"),
    BotCommand(command="role", description="Модель поведения"),
    BotCommand(command="status", description="Статус памяти"),
    BotCommand(command="clear_chat", description="Очистить историю диалога"),
    BotCommand(command="clear_docs", description="Удалить документы"),
    BotCommand(command="clear_all", description="Очистить всё"),
    BotCommand(command="help", description="Подробная справка"),
]

# Тексты кнопок инлайн-меню (сетка над полем ввода)
BTN_START = "🚀 Начать"
BTN_ROLE = "🎭 Роль"
BTN_STATUS = "📊 Статус"
BTN_CLEAR_CHAT = "🧹 Очистить чат"
BTN_CLEAR_DOCS = "📄 Удалить док."
BTN_CLEAR_ALL = "🗑️ Очистить всё"
BTN_HELP = "❓ Справка"

# Клавиатура-меню (сетка кнопок над полем ввода, как во вложении)
menu_keyboard = ReplyKeyboardMarkup(
    keyboard=[
        [KeyboardButton(text=BTN_START)],
        [KeyboardButton(text=BTN_ROLE), KeyboardButton(text=BTN_STATUS)],
        [KeyboardButton(text=BTN_CLEAR_CHAT), KeyboardButton(text=BTN_CLEAR_DOCS)],
        [KeyboardButton(text=BTN_CLEAR_ALL)],
        [KeyboardButton(text=BTN_HELP)],
    ],
    resize_keyboard=True,
)


async def _do_start(message: Message):
    """Общая логика «Начать работу»."""
    user_id = message.from_user.id
    if user_id in user_histories:
        user_histories[user_id].clear()
    await message.answer(
        "👋 <b>Привет! Я бот с гибридной памятью.</b>\n\n"
        "💭 <b>Короткая память:</b> последние 10 сообщений диалога\n"
        "📚 <b>Долгая память:</b> ваши документы (PDF, TXT, DOCX)\n\n"
        "Давай начнём: можешь выбрать роль в меню ниже, загрузить документ или просто напиши мне.",
        parse_mode="HTML",
        reply_markup=menu_keyboard,
    )


@dp.message(Command("start"))
@dp.message(F.text == BTN_START)
async def cmd_start(message: Message):
    await _do_start(message)


def _role_keyboard() -> InlineKeyboardMarkup:
    """Инлайн-кнопки выбора роли (5 ролей)."""
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=ROLES["standard"][0], callback_data="role:standard")],
        [InlineKeyboardButton(text=ROLES["scientific"][0], callback_data="role:scientific")],
        [InlineKeyboardButton(text=ROLES["technical"][0], callback_data="role:technical")],
        [InlineKeyboardButton(text=ROLES["creative"][0], callback_data="role:creative")],
        [InlineKeyboardButton(text=ROLES["concise"][0], callback_data="role:concise")],
    ])


@dp.message(Command("role"))
@dp.message(F.text == BTN_ROLE)
async def cmd_role(message: Message):
    """Модель поведения: текущая роль и выбор из 5 ролей."""
    user_id = message.from_user.id
    current = get_user_role(user_id)
    name, _ = ROLES.get(current, ROLES[DEFAULT_ROLE])
    await message.answer(
        "🎭 <b>Модель поведения</b>\n\n"
        f"Текущая роль: <b>{name}</b>\n\n"
        "Выбери одну из ролей — от неё зависит стиль и тон ответов. "
        "Память (диалог, документы, тезисы) работает во всех ролях.\n\n"
        f"Модель: <code>{OPENAI_MODEL}</code>",
        parse_mode="HTML",
        reply_markup=_role_keyboard(),
    )


@dp.callback_query(F.data.startswith("role:"))
async def cb_role(callback: CallbackQuery):
    """Обработка нажатия кнопки выбора роли."""
    role_key = callback.data.removeprefix("role:")
    if role_key not in ROLES:
        await callback.answer("Неизвестная роль")
        return
    user_id = callback.from_user.id
    user_roles[user_id] = role_key
    name = ROLES[role_key][0]
    await callback.message.edit_text(
        f"🎭 <b>Модель поведения</b>\n\n"
        f"Роль установлена: <b>{name}</b>\n\n"
        "Следующие ответы будут в выбранном стиле.\n\n"
        f"Модель: <code>{OPENAI_MODEL}</code>",
        parse_mode="HTML",
    )
    await callback.answer(f"Роль: {name}")


@dp.message(Command("help"))
@dp.message(F.text == BTN_HELP)
async def cmd_help(message: Message):
    """Подробная справка"""
    await message.answer(
        "📚 <b>Подробная справка</b>\n\n"
        "<b>🎯 Как использовать:</b>\n\n"
        "1️⃣ <b>Диалог</b> — пишите, я помню последние 10 сообщений\n\n"
        "2️⃣ <b>Документы</b> — отправьте PDF/TXT/DOCX, потом задавайте вопросы по ним\n\n"
        "3️⃣ <b>Гибрид</b> — документы + история диалога подставляются автоматически\n\n"
        "<b>📋 Команды:</b>\n"
        "/start — Начать работу\n"
        "/role — Модель поведения\n"
        "/status — Статус памяти\n"
        "/clear_chat — Очистить историю диалога\n"
        "/clear_docs — Удалить документы\n"
        "/clear_all — Очистить всё",
        parse_mode="HTML"
    )


@dp.message(Command("status"))
@dp.message(F.text == BTN_STATUS)
async def cmd_status(message: Message):
    """Показать статус обеих памятей"""
    user_id = message.from_user.id
    
    # Короткая память
    history = get_user_history(user_id)
    chat_messages = len(history)
    
    # Долгая память
    docs_count = get_user_documents_count(user_id)
    total_docs = collection.count()
    
    status_text = (
        "📊 <b>Статус памяти</b>\n\n"
        "💭 <b>Короткая память (диалог):</b>\n"
        f"  • Сообщений в истории: <b>{chat_messages}/{HISTORY_SIZE}</b>\n\n"
        "📚 <b>Долгая память (документы):</b>\n"
        f"  • Ваших фрагментов: <b>{docs_count}</b>\n"
        f"  • Всего в базе: <b>{total_docs}</b>\n\n"
        "⚙️ <b>Настройки:</b>\n"
        f"  • Модель: <code>{OPENAI_MODEL}</code>\n"
        f"  • Эмбеддинги: <code>{EMBEDDING_MODEL}</code>\n"
        f"  • База: <code>{MEMORY_PATH}</code>"
    )
    
    await message.answer(status_text, parse_mode="HTML")


@dp.message(Command("clear_chat"))
@dp.message(F.text == BTN_CLEAR_CHAT)
async def cmd_clear_chat(message: Message):
    """Очистить короткую память (историю диалога)"""
    user_id = message.from_user.id
    
    if user_id in user_histories:
        user_histories[user_id].clear()
    
    await message.answer(
        "🧹 <b>История диалога очищена!</b>\n\n"
        "Короткая память сброшена.\n"
        "Документы остались на месте.",
        parse_mode="HTML"
    )


@dp.message(Command("clear_docs"))
@dp.message(F.text == BTN_CLEAR_DOCS)
async def cmd_clear_docs(message: Message):
    """Очистить долгую память (документы)"""
    user_id = message.from_user.id
    deleted_count = delete_user_documents(user_id)
    
    if deleted_count > 0:
        await message.answer(
            f"🗑️ <b>Документы удалены!</b>\n\n"
            f"Удалено фрагментов: <b>{deleted_count}</b>\n"
            f"История диалога сохранена.",
            parse_mode="HTML"
        )
    else:
        await message.answer(
            "📭 <b>Документов нет</b>\n\n"
            "У вас нет загруженных документов.",
            parse_mode="HTML"
        )


@dp.message(Command("clear_all"))
@dp.message(F.text == BTN_CLEAR_ALL)
async def cmd_clear_all(message: Message):
    """Очистить всё (и короткую, и долгую память)"""
    user_id = message.from_user.id
    
    # Очищаем короткую память
    if user_id in user_histories:
        user_histories[user_id].clear()
    
    # Очищаем долгую память
    deleted_count = delete_user_documents(user_id)
    
    await message.answer(
        "🧹 <b>Вся память очищена!</b>\n\n"
        f"• История диалога: сброшена\n"
        f"• Документы: удалено {deleted_count} фрагментов\n\n"
        "Можно начинать заново!",
        parse_mode="HTML"
    )


# ============================================
# ОБРАБОТЧИКИ ДОКУМЕНТОВ И СООБЩЕНИЙ
# ============================================

@dp.message(F.document)
async def handle_document(message: Message):
    """Обработчик загрузки документов"""
    user_id = message.from_user.id
    document: Document = message.document
    file_name = document.file_name
    file_extension = Path(file_name).suffix.lower()
    
    if file_extension not in ['.pdf', '.txt', '.docx']:
        await message.answer(
            f"❌ Неподдерживаемый формат: <code>{file_extension}</code>\n\n"
            f"Поддерживаются: PDF, TXT, DOCX",
            parse_mode="HTML"
        )
        return
    
    try:
        status_msg = await message.answer(
            f"⏳ Обрабатываю документ <b>{file_name}</b>...",
            parse_mode="HTML"
        )
        
        # Скачиваем файл
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            tmp_path = tmp_file.name
            await bot.download(document, destination=tmp_path)
        
        # Извлекаем текст
        text = load_document(tmp_path, file_extension)
        os.unlink(tmp_path)
        
        if len(text.strip()) < 50:
            await status_msg.edit_text(
                "❌ Документ слишком короткий или пустой."
            )
            return
        
        # Разбиваем на чанки
        chunks = split_text_into_chunks(text)
        
        await status_msg.edit_text(
            f"🔄 Создание эмбеддингов для {len(chunks)} фрагментов..."
        )
        
        # Сохраняем в долгую память
        saved_count = await embed_chunks(user_id, file_name, chunks)
        
        await status_msg.edit_text(
            f"✅ <b>Документ сохранён в долгую память!</b>\n\n"
            f"📄 Файл: <code>{file_name}</code>\n"
            f"📊 Символов: <b>{len(text)}</b>\n"
            f"🗂️ Фрагментов: <b>{saved_count}</b>\n\n"
            f"Теперь я могу отвечать на вопросы по этому документу!",
            parse_mode="HTML"
        )
        
    except Exception as e:
        logger.error(f"Ошибка обработки документа: {e}")
        await message.answer(
            f"❌ Ошибка обработки документа:\n<code>{str(e)}</code>",
            parse_mode="HTML"
        )


@dp.message(F.text)
async def handle_text_message(message: Message):
    """Обработчик текстовых сообщений"""
    user_id = message.from_user.id
    user_text = message.text
    
    logger.info(f"Сообщение от user_id={user_id}: {user_text[:100]}")
    database.add_user_request(user_id, user_text)
    export_long_memory_to_json()

    # Показываем индикатор печати
    await message.bot.send_chat_action(
        chat_id=message.chat.id,
        action="typing"
    )
    
    # Получаем гибридный ответ (короткая + долгая память)
    ai_response = await get_hybrid_response(user_id, user_text)
    
    # Отправляем ответ
    await message.answer(ai_response)


# ============================================
# ГЛАВНАЯ ФУНКЦИЯ
# ============================================

async def main():
    """Главная функция запуска бота"""
    logger.info("🚀 Бот с гибридной памятью запущен!")
    logger.info(f"💭 Короткая память: {HISTORY_SIZE} сообщений")
    logger.info(f"📚 Долгая память: {collection.count()} фрагментов в базе")
    export_long_memory_to_json()
    
    await bot.delete_webhook(drop_pending_updates=True)
    await bot.set_my_commands(BOT_COMMANDS)
    await dp.start_polling(bot)


if __name__ == "__main__":
    import asyncio
    
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("⏹️ Бот остановлен пользователем")

