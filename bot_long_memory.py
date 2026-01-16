"""
Telegram-бот с долгой памятью (RAG на основе ChromaDB)
Использует aiogram 3.x, OpenAI Embeddings и ChromaDB
"""

import os
import logging
import tempfile
from typing import List, Dict, Optional
from pathlib import Path

from aiogram import Bot, Dispatcher, F
from aiogram.filters import Command
from aiogram.types import Message, Document
from openai import AsyncOpenAI
from dotenv import load_dotenv

import chromadb
from chromadb.config import Settings

# Библиотеки для работы с документами
import PyPDF2
from docx import Document as DocxDocument


# Загрузка переменных окружения
load_dotenv()

# Загрузка настроек
BOT_TOKEN = os.getenv("BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "").strip()
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5-mini")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-3-small")
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()

# Настройка логирования
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Проверка обязательных параметров
if not BOT_TOKEN:
    raise ValueError("BOT_TOKEN не найден в переменных окружения!")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY не найден в переменных окружения!")

logger.info(f"Бот инициализирован. Модель: {OPENAI_MODEL}, Эмбеддинги: {EMBEDDING_MODEL}")

# Инициализация бота и диспетчера
bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

# Инициализация OpenAI клиента
if OPENAI_BASE_URL:
    openai_client = AsyncOpenAI(
        api_key=OPENAI_API_KEY,
        base_url=OPENAI_BASE_URL
    )
else:
    openai_client = AsyncOpenAI(api_key=OPENAI_API_KEY)

# Инициализация ChromaDB (persistent storage)
MEMORY_PATH = "./memory"
Path(MEMORY_PATH).mkdir(exist_ok=True)

chroma_client = chromadb.PersistentClient(
    path=MEMORY_PATH,
    settings=Settings(anonymized_telemetry=False)
)

# Коллекция для хранения эмбеддингов (создаётся если не существует)
try:
    collection = chroma_client.get_or_create_collection(
        name="documents",
        metadata={"description": "User documents with embeddings"}
    )
    logger.info(f"ChromaDB коллекция готова. Документов в базе: {collection.count()}")
except Exception as e:
    logger.error(f"Ошибка инициализации ChromaDB: {e}")
    raise

# Константы
CHUNK_SIZE = 500  # Размер чанка в символах
CHUNK_OVERLAP = 50  # Перекрытие между чанками


# ============================================
# ФУНКЦИИ ДЛЯ РАБОТЫ С ДОКУМЕНТАМИ
# ============================================

def load_document(file_path: str, file_extension: str) -> str:
    """
    Загрузка и извлечение текста из документа.
    
    Args:
        file_path: Путь к файлу
        file_extension: Расширение файла (.pdf, .txt, .docx)
    
    Returns:
        Извлечённый текст
    """
    try:
        if file_extension == '.txt':
            # Обработка TXT
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            logger.info(f"TXT загружен, длина: {len(text)} символов")
            return text
        
        elif file_extension == '.pdf':
            # Обработка PDF
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
            logger.info(f"PDF загружен ({len(pdf_reader.pages)} страниц), длина: {len(text)} символов")
            return text
        
        elif file_extension == '.docx':
            # Обработка DOCX
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            logger.info(f"DOCX загружен ({len(doc.paragraphs)} параграфов), длина: {len(text)} символов")
            return text
        
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {file_extension}")
    
    except Exception as e:
        logger.error(f"Ошибка загрузки документа: {e}")
        raise


def split_text_into_chunks(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    Разбиение текста на чанки с перекрытием.
    
    Args:
        text: Исходный текст
        chunk_size: Размер чанка в символах
        overlap: Размер перекрытия между чанками
    
    Returns:
        Список чанков
    """
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]
        
        # Пропускаем пустые чанки
        if chunk.strip():
            chunks.append(chunk.strip())
        
        start += chunk_size - overlap
    
    logger.info(f"Текст разбит на {len(chunks)} чанков")
    return chunks


async def embed_chunks(user_id: int, document_name: str, chunks: List[str]) -> int:
    """
    Создание эмбеддингов для чанков и сохранение в ChromaDB.
    
    Args:
        user_id: ID пользователя
        document_name: Название документа
        chunks: Список чанков текста
    
    Returns:
        Количество сохранённых чанков
    """
    try:
        logger.info(f"Создание эмбеддингов для {len(chunks)} чанков...")
        
        # Создаём эмбеддинги через OpenAI API
        embeddings_response = await openai_client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=chunks
        )
        
        # Извлекаем векторы эмбеддингов
        embeddings = [item.embedding for item in embeddings_response.data]
        
        # Генерируем уникальные ID для чанков
        ids = [f"user_{user_id}_doc_{document_name}_chunk_{i}" for i in range(len(chunks))]
        
        # Метаданные для каждого чанка
        metadatas = [
            {
                "user_id": str(user_id),
                "document_name": document_name,
                "chunk_index": i,
                "chunk_text": chunks[i][:100]  # Первые 100 символов для предпросмотра
            }
            for i in range(len(chunks))
        ]
        
        # Сохраняем в ChromaDB
        collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas
        )
        
        logger.info(f"✅ Сохранено {len(chunks)} чанков в ChromaDB")
        return len(chunks)
    
    except Exception as e:
        logger.error(f"Ошибка создания эмбеддингов: {e}")
        raise


async def retrieve_context(user_id: int, query: str, top_k: int = 3) -> List[str]:
    """
    Поиск релевантных фрагментов в векторной базе.
    
    Args:
        user_id: ID пользователя
        query: Запрос пользователя
        top_k: Количество наиболее релевантных фрагментов
    
    Returns:
        Список релевантных текстовых фрагментов
    """
    try:
        logger.info(f"Поиск контекста для запроса: '{query[:50]}...'")
        
        # Создаём эмбеддинг для запроса
        query_embedding_response = await openai_client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=[query]
        )
        query_embedding = query_embedding_response.data[0].embedding
        
        # Поиск в ChromaDB с фильтром по user_id
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            where={"user_id": str(user_id)}
        )
        
        # Извлекаем найденные документы
        if results and results['documents'] and len(results['documents']) > 0:
            documents = results['documents'][0]
            logger.info(f"Найдено {len(documents)} релевантных фрагментов")
            return documents
        else:
            logger.warning("Релевантные фрагменты не найдены")
            return []
    
    except Exception as e:
        logger.error(f"Ошибка поиска контекста: {e}")
        return []


async def answer_question(user_id: int, question: str) -> str:
    """
    Генерация ответа на вопрос на основе документов пользователя.
    
    Args:
        user_id: ID пользователя
        question: Вопрос пользователя
    
    Returns:
        Ответ модели
    """
    try:
        # Получаем релевантный контекст из векторной базы
        context_chunks = await retrieve_context(user_id, question, top_k=3)
        
        if not context_chunks:
            return (
                "❌ Не удалось найти релевантную информацию в ваших документах.\n\n"
                "Убедитесь, что вы загрузили документ командой /upload или отправкой файла."
            )
        
        # Формируем контекст из найденных фрагментов
        context = "\n\n---\n\n".join(context_chunks)
        
        # Формируем промпт для модели
        system_prompt = (
            "Ты — AI-ассистент, который отвечает на вопросы ТОЛЬКО на основе предоставленного контекста из документов пользователя.\n"
            "Правила:\n"
            "1. Отвечай только на основе контекста\n"
            "2. Если ответа нет в контексте — так и скажи\n"
            "3. Не выдумывай информацию\n"
            "4. Будь точным и кратким\n"
            "5. Цитируй документ, если это уместно"
        )
        
        user_prompt = f"""Контекст из документов:
{context}

Вопрос пользователя: {question}

Ответ:"""
        
        # Запрос к OpenAI API
        logger.info(f"Генерация ответа для user_id={user_id}")
        
        # Для gpt-5-mini не используем temperature (только значение по умолчанию 1)
        response = await openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_completion_tokens=1000  # Для gpt-5-mini используется max_completion_tokens
        )
        
        answer = response.choices[0].message.content
        logger.info(f"Ответ сгенерирован, длина: {len(answer)} символов")
        
        return answer
    
    except Exception as e:
        logger.error(f"Ошибка генерации ответа: {e}")
        return f"❌ Произошла ошибка при генерации ответа: {str(e)}"


def get_user_documents_count(user_id: int) -> int:
    """
    Получить количество чанков пользователя в базе.
    
    Args:
        user_id: ID пользователя
    
    Returns:
        Количество чанков
    """
    try:
        results = collection.get(
            where={"user_id": str(user_id)}
        )
        return len(results['ids']) if results and results['ids'] else 0
    except Exception as e:
        logger.error(f"Ошибка подсчёта документов: {e}")
        return 0


def delete_user_documents(user_id: int) -> int:
    """
    Удалить все документы пользователя из базы.
    
    Args:
        user_id: ID пользователя
    
    Returns:
        Количество удалённых чанков
    """
    try:
        # Получаем все ID документов пользователя
        results = collection.get(
            where={"user_id": str(user_id)}
        )
        
        if results and results['ids']:
            collection.delete(ids=results['ids'])
            logger.info(f"Удалено {len(results['ids'])} чанков для user_id={user_id}")
            return len(results['ids'])
        return 0
    except Exception as e:
        logger.error(f"Ошибка удаления документов: {e}")
        return 0


# ============================================
# ОБРАБОТЧИКИ КОМАНД
# ============================================

@dp.message(Command("start"))
async def cmd_start(message: Message):
    """
    Обработчик команды /start
    """
    await message.answer(
        "👋 <b>Привет! Я бот с долгой памятью.</b>\n\n"
        "Я могу запоминать ваши документы и отвечать на вопросы по ним!\n\n"
        "📄 <b>Поддерживаемые форматы:</b>\n"
        "• PDF (.pdf)\n"
        "• Текстовые файлы (.txt)\n"
        "• Word документы (.docx)\n\n"
        "📋 <b>Как использовать:</b>\n"
        "1. Отправьте мне документ (просто прикрепите файл)\n"
        "2. Я обработаю его и сохраню в память\n"
        "3. Задавайте вопросы — я отвечу на основе ваших документов\n\n"
        "⚙️ <b>Команды:</b>\n"
        "/start - Показать это сообщение\n"
        "/status - Статус вашей базы знаний\n"
        "/clear - Очистить базу знаний\n"
        "/help - Подробная справка",
        parse_mode="HTML"
    )


@dp.message(Command("help"))
async def cmd_help(message: Message):
    """
    Обработчик команды /help
    """
    await message.answer(
        "📚 <b>Подробная справка</b>\n\n"
        "<b>Как это работает:</b>\n"
        "1. Вы загружаете документ\n"
        "2. Бот разбивает его на фрагменты (~500 символов)\n"
        "3. Создаются эмбеддинги (векторные представления текста)\n"
        "4. Сохраняется в ChromaDB (векторная база данных)\n"
        "5. При вопросе бот ищет релевантные фрагменты\n"
        "6. AI генерирует ответ на основе найденной информации\n\n"
        "<b>Преимущества:</b>\n"
        "✅ Долговременная память (данные сохраняются)\n"
        "✅ Точные ответы (только из ваших документов)\n"
        "✅ Поддержка больших документов\n"
        "✅ Семантический поиск (понимает смысл)\n\n"
        "<b>Советы:</b>\n"
        "• Загружайте документы с чёткой структурой\n"
        "• Задавайте конкретные вопросы\n"
        "• Используйте /status для проверки базы\n"
        "• Очищайте базу /clear если документы устарели",
        parse_mode="HTML"
    )


@dp.message(Command("status"))
async def cmd_status(message: Message):
    """
    Обработчик команды /status - показать статус базы знаний
    """
    user_id = message.from_user.id
    chunks_count = get_user_documents_count(user_id)
    total_in_db = collection.count()
    
    if chunks_count == 0:
        await message.answer(
            "📊 <b>Статус базы знаний</b>\n\n"
            "❌ У вас пока нет загруженных документов.\n\n"
            "Отправьте мне документ (PDF, TXT или DOCX), чтобы начать!",
            parse_mode="HTML"
        )
    else:
        await message.answer(
            f"📊 <b>Статус базы знаний</b>\n\n"
            f"✅ Ваших фрагментов в базе: <b>{chunks_count}</b>\n"
            f"📚 Всего в базе: <b>{total_in_db}</b>\n"
            f"💾 База данных: <code>{MEMORY_PATH}</code>\n"
            f"🤖 Модель: <code>{OPENAI_MODEL}</code>\n"
            f"🧮 Эмбеддинги: <code>{EMBEDDING_MODEL}</code>\n\n"
            f"Задавайте вопросы — я отвечу на основе ваших документов!",
            parse_mode="HTML"
        )


@dp.message(Command("clear"))
async def cmd_clear(message: Message):
    """
    Обработчик команды /clear - очистить базу знаний пользователя
    """
    user_id = message.from_user.id
    deleted_count = delete_user_documents(user_id)
    
    if deleted_count > 0:
        await message.answer(
            f"🧹 <b>База знаний очищена!</b>\n\n"
            f"Удалено фрагментов: <b>{deleted_count}</b>\n\n"
            f"Теперь вы можете загрузить новые документы.",
            parse_mode="HTML"
        )
    else:
        await message.answer(
            "📭 <b>База знаний пуста</b>\n\n"
            "У вас нет загруженных документов для удаления.",
            parse_mode="HTML"
        )


# ============================================
# ОБРАБОТЧИКИ ДОКУМЕНТОВ И СООБЩЕНИЙ
# ============================================

@dp.message(F.document)
async def handle_document(message: Message):
    """
    Обработчик загрузки документов
    """
    user_id = message.from_user.id
    document: Document = message.document
    file_name = document.file_name
    file_extension = Path(file_name).suffix.lower()
    
    # Проверка поддерживаемых форматов
    if file_extension not in ['.pdf', '.txt', '.docx']:
        await message.answer(
            f"❌ Неподдерживаемый формат: <code>{file_extension}</code>\n\n"
            f"Поддерживаются: PDF, TXT, DOCX",
            parse_mode="HTML"
        )
        return
    
    try:
        # Отправляем статус обработки
        status_msg = await message.answer(
            f"⏳ Обрабатываю документ <b>{file_name}</b>...\n"
            f"Это может занять некоторое время.",
            parse_mode="HTML"
        )
        
        # Скачиваем файл во временную папку
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            tmp_path = tmp_file.name
            await bot.download(document, destination=tmp_path)
        
        logger.info(f"Документ {file_name} скачан для user_id={user_id}")
        
        # Загружаем и извлекаем текст
        text = load_document(tmp_path, file_extension)
        
        # Удаляем временный файл
        os.unlink(tmp_path)
        
        if len(text.strip()) < 50:
            await status_msg.edit_text(
                "❌ Документ слишком короткий или пустой.\n"
                "Убедитесь, что файл содержит текст."
            )
            return
        
        # Разбиваем на чанки
        chunks = split_text_into_chunks(text)
        
        # Создаём эмбеддинги и сохраняем
        await status_msg.edit_text(
            f"🔄 Создание эмбеддингов для {len(chunks)} фрагментов..."
        )
        
        saved_count = await embed_chunks(user_id, file_name, chunks)
        
        # Успешное завершение
        await status_msg.edit_text(
            f"✅ <b>Документ обработан!</b>\n\n"
            f"📄 Файл: <code>{file_name}</code>\n"
            f"📊 Символов: <b>{len(text)}</b>\n"
            f"🗂️ Фрагментов: <b>{saved_count}</b>\n\n"
            f"Теперь вы можете задавать вопросы по этому документу!",
            parse_mode="HTML"
        )
        
    except Exception as e:
        logger.error(f"Ошибка обработки документа: {e}")
        await message.answer(
            f"❌ Ошибка при обработке документа:\n"
            f"<code>{str(e)}</code>\n\n"
            f"Попробуйте загрузить другой файл.",
            parse_mode="HTML"
        )


@dp.message(F.text)
async def handle_text_message(message: Message):
    """
    Обработчик текстовых сообщений (вопросов)
    """
    user_id = message.from_user.id
    question = message.text
    
    # Проверяем, есть ли у пользователя документы
    chunks_count = get_user_documents_count(user_id)
    if chunks_count == 0:
        await message.answer(
            "📭 <b>У вас нет загруженных документов</b>\n\n"
            "Сначала загрузите документ, затем задавайте вопросы.\n"
            "Отправьте мне файл (PDF, TXT или DOCX).",
            parse_mode="HTML"
        )
        return
    
    # Показываем индикатор печати
    await message.bot.send_chat_action(
        chat_id=message.chat.id,
        action="typing"
    )
    
    logger.info(f"Вопрос от user_id={user_id}: {question[:100]}")
    
    # Генерируем ответ
    answer = await answer_question(user_id, question)
    
    # Отправляем ответ
    await message.answer(answer)


# ============================================
# ГЛАВНАЯ ФУНКЦИЯ
# ============================================

async def main():
    """
    Главная функция запуска бота
    """
    logger.info("🚀 Бот с долгой памятью запущен!")
    logger.info(f"📁 Путь к базе данных: {MEMORY_PATH}")
    logger.info(f"📊 Документов в базе: {collection.count()}")
    
    # Удаляем webhook если был установлен
    await bot.delete_webhook(drop_pending_updates=True)
    
    # Запускаем polling
    await dp.start_polling(bot)


if __name__ == "__main__":
    import asyncio
    
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("⏹️ Бот остановлен пользователем")

